"""
README.md → Professional README.docx converter
Features: Cover page, TOC, styled headings, code blocks, tables,
          bullet/numbered lists, inline formatting, page numbers, footer
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

README = "project_report.md"
OUTPUT = "project_report.docx"

doc = Document()

# ── Page Setup ────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)

# ── Color Palette ─────────────────────────────────────────────────────────────
COLOR_PRIMARY    = RGBColor(0x1E, 0x3A, 0x5F)   # Deep navy
COLOR_H2         = RGBColor(0x2E, 0x6D, 0xA4)   # Steel blue
COLOR_H3         = RGBColor(0x35, 0x7A, 0x6E)   # Teal
COLOR_H4         = RGBColor(0x5A, 0x5A, 0x5A)   # Dark grey
COLOR_CODE_FG    = RGBColor(0x1F, 0x2D, 0x3D)   # Near black
COLOR_INLINE_FG  = RGBColor(0xC0, 0x39, 0x2B)   # Deep red
COLOR_ACCENT     = RGBColor(0x1E, 0x3A, 0x5F)   # Navy accent
FILL_CODE        = "F0F4F8"                       # Light blue-grey
FILL_COVER_TOP   = "1E3A5F"                       # Navy
FILL_COVER_BAND  = "2E6DA4"                       # Blue band

# ── XML Helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)

def add_border_bottom(para, color="CCCCCC", sz="4"):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)

def add_page_number(para):
    """Insert auto page number field into paragraph."""
    run = para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run2 = para.add_run()
    ins  = OxmlElement("w:instrText")
    ins.set(qn("xml:space"), "preserve")
    ins.text = " PAGE "
    run2._r.append(ins)

    run3 = para.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)

def add_page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

# ── Inline Formatter ──────────────────────────────────────────────────────────
def add_inline(para, text, base_size=10.5):
    """Parse **bold**, *italic*, `inline code` and add styled runs."""
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts   = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name  = "Consolas"
            run.font.size  = Pt(base_size - 1)
            run.font.color.rgb = COLOR_INLINE_FG
            # light highlight
            rPr = run._r.get_or_add_rPr()
            hl  = OxmlElement("w:highlight")
            hl.set(qn("w:val"), "yellow")
            rPr.append(hl)
        else:
            run = para.add_run(part)
            run.font.size = Pt(base_size)

# ── Cover Page ────────────────────────────────────────────────────────────────
def add_cover_page(doc, title, subtitle):
    # Top navy block
    top = doc.add_paragraph()
    top.paragraph_format.space_before = Pt(0)
    top.paragraph_format.space_after  = Pt(0)
    set_para_bg(top, FILL_COVER_TOP)
    run = top.add_run(" " * 200)
    run.font.size = Pt(72)

    # Title block
    for _ in range(6):
        sp = doc.add_paragraph()
        set_para_bg(sp, FILL_COVER_TOP)
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_bg(t, FILL_COVER_TOP)
    r = t.add_run(title)
    r.font.size  = Pt(28)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.name  = "Calibri"
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after  = Pt(6)

    # Blue band subtitle
    band = doc.add_paragraph()
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_bg(band, FILL_COVER_BAND)
    rb = band.add_run(subtitle)
    rb.font.size  = Pt(13)
    rb.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rb.font.name  = "Calibri"
    rb.font.italic = True
    band.paragraph_format.space_before = Pt(4)
    band.paragraph_format.space_after  = Pt(4)

    # Spacer rows
    for _ in range(4):
        sp = doc.add_paragraph()
        set_para_bg(sp, FILL_COVER_TOP)
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    # Meta info
    meta_items = [
        ("Document Type",  "Technical Project Report"),
        ("Version",        "2.0.0"),
        ("Date",           datetime.now().strftime("%B %d, %Y")),
        ("Technology",     "FastAPI · Groq LLaMA 3.3-70B · FAISS · Neo4j · Redis"),
    ]
    for label, value in meta_items:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_bg(mp, FILL_COVER_TOP)
        rl = mp.add_run(f"{label}:  ")
        rl.font.size  = Pt(10)
        rl.font.bold  = True
        rl.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
        rl.font.name  = "Calibri"
        rv = mp.add_run(value)
        rv.font.size  = Pt(10)
        rv.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rv.font.name  = "Calibri"
        mp.paragraph_format.space_before = Pt(2)
        mp.paragraph_format.space_after  = Pt(2)

    # Bottom spacer
    for _ in range(3):
        sp = doc.add_paragraph()
        set_para_bg(sp, FILL_COVER_TOP)
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    add_page_break(doc)

# ── Footer with Page Numbers ──────────────────────────────────────────────────
def add_footer(doc, project_name):
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_border_bottom(fp, color="2E6DA4", sz="6")

        left = fp.add_run(f"  {project_name}  |  Technical Report  |  ")
        left.font.size  = Pt(8)
        left.font.color.rgb = COLOR_H4
        left.font.name  = "Calibri"

        add_page_number(fp)

        right = fp.add_run(f"  |  {datetime.now().strftime('%Y')}  ")
        right.font.size  = Pt(8)
        right.font.color.rgb = COLOR_H4
        right.font.name  = "Calibri"

# ── Heading Styler ────────────────────────────────────────────────────────────
def style_heading(para, level):
    sizes  = {1: 20, 2: 16, 3: 13, 4: 11}
    colors = {1: COLOR_PRIMARY, 2: COLOR_H2, 3: COLOR_H3, 4: COLOR_H4}
    names  = {1: "Calibri", 2: "Calibri", 3: "Calibri", 4: "Calibri"}

    for run in para.runs:
        run.bold           = True
        run.font.size      = Pt(sizes.get(level, 11))
        run.font.color.rgb = colors.get(level, COLOR_H4)
        run.font.name      = names.get(level, "Calibri")

    para.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10, 4: 8}.get(level, 8))
    para.paragraph_format.space_after  = Pt({1: 8,  2: 6,  3: 4,  4: 3}.get(level, 3))

    if level in (1, 2):
        add_border_bottom(para, color={1: "1E3A5F", 2: "2E6DA4"}.get(level), sz="8" if level == 1 else "4")

# ── Code Block Styler ─────────────────────────────────────────────────────────
def render_code_block(doc, lines):
    # Top label bar
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(8)
    label.paragraph_format.space_after  = Pt(0)
    set_para_bg(label, "2E6DA4")
    lr = label.add_run("  CODE")
    lr.font.name  = "Consolas"
    lr.font.size  = Pt(7.5)
    lr.font.bold  = True
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Code lines
    for cl in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)
        set_para_bg(p, FILL_CODE)
        run = p.add_run(cl if cl else " ")
        run.font.name  = "Consolas"
        run.font.size  = Pt(8.5)
        run.font.color.rgb = COLOR_CODE_FG

    # Bottom border line
    bot = doc.add_paragraph()
    bot.paragraph_format.space_before = Pt(0)
    bot.paragraph_format.space_after  = Pt(8)
    set_para_bg(bot, "D0DCE8")
    bot.add_run(" ")

# ── Table Renderer ────────────────────────────────────────────────────────────
def render_table(doc, table_buf):
    rows = [r for r in table_buf if not re.match(r'^\|[-| :]+\|$', r.strip())]
    if not rows:
        return
    parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    ncols  = max(len(r) for r in parsed)

    t = doc.add_table(rows=len(parsed), cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(parsed):
        for ci in range(ncols):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            text = row[ci] if ci < len(row) else ""

            if ri == 0:
                set_cell_bg(cell, "1E3A5F")
                run = p.add_run(text)
                run.bold           = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size      = Pt(9.5)
                run.font.name      = "Calibri"
                p.alignment        = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = "EEF3F8" if ri % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                add_inline(p, text, base_size=9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN PARSE LOOP
# ══════════════════════════════════════════════════════════════════════════════
with open(README, encoding="utf-8") as f:
    lines = f.readlines()

# Extract title + subtitle for cover
cover_title    = "AI Research Explainer Engine"
cover_subtitle = "AI-Powered Research Intelligence & Multi-Level Explanation System"
for line in lines[:5]:
    if line.startswith("# "):
        cover_title = line[2:].strip()
    if line.startswith("**") and line.endswith("**\n"):
        cover_subtitle = line.strip().strip("*")

add_cover_page(doc, cover_title, cover_subtitle)
add_footer(doc, cover_title)

in_code   = False
code_buf  = []
table_buf = []

i = 0
while i < len(lines):
    line = lines[i].rstrip("\n")

    # ── Skip the H1 title and bold subtitle (already on cover) ───────────────
    if line.startswith("# ") and i < 3:
        i += 1
        continue
    if line.startswith("**") and line.endswith("**") and i < 5:
        i += 1
        continue

    # ── Code fence ────────────────────────────────────────────────────────────
    if line.strip().startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            render_code_block(doc, code_buf)
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # ── Table ─────────────────────────────────────────────────────────────────
    if line.startswith("|"):
        table_buf.append(line)
        i += 1
        continue
    else:
        if table_buf:
            render_table(doc, table_buf)
            table_buf = []

    # ── Horizontal rule ───────────────────────────────────────────────────────
    if re.match(r'^---+$', line.strip()):
        p = doc.add_paragraph()
        add_border_bottom(p, color="2E6DA4", sz="4")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        i += 1
        continue

    # ── Headings ──────────────────────────────────────────────────────────────
    hm = re.match(r'^(#{1,4})\s+(.*)', line)
    if hm:
        level = len(hm.group(1))
        text  = hm.group(2).strip()
        p = doc.add_heading("", level=level)
        p.clear()
        add_inline(p, text)
        style_heading(p, level)
        i += 1
        continue

    # ── Italic figure caption (*Figure N: ...*) ───────────────────────────────
    fig = re.match(r'^\*(.+)\*$', line.strip())
    if fig:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(fig.group(1))
        r.italic         = True
        r.font.size      = Pt(9)
        r.font.color.rgb = COLOR_H4
        r.font.name      = "Calibri"
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(8)
        i += 1
        continue

    # ── Bullet list ───────────────────────────────────────────────────────────
    bm = re.match(r'^(\s*)[-*]\s+(.*)', line)
    if bm:
        indent = len(bm.group(1)) // 2
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent  = Inches(0.3 * (indent + 1))
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        add_inline(p, bm.group(2))
        i += 1
        continue

    # ── Numbered list ─────────────────────────────────────────────────────────
    nm = re.match(r'^\s*\d+\.\s+(.*)', line)
    if nm:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        add_inline(p, nm.group(1))
        i += 1
        continue

    # ── Blank line ────────────────────────────────────────────────────────────
    if line.strip() == "":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        i += 1
        continue

    # ── Normal paragraph ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    add_inline(p, line)
    i += 1

doc.save(OUTPUT)
print(f"✅  Saved → {OUTPUT}")
